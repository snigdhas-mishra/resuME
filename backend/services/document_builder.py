"""
Document builder service — converts Markdown resume text into DOCX format.
"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_styled_paragraph(doc: Document, text: str, style_name: str, **kwargs):
    """Add a paragraph with a given style and optional formatting overrides."""
    para = doc.add_paragraph()

    # Apply font size
    font_size = kwargs.get("font_size")
    bold = kwargs.get("bold", False)
    color = kwargs.get("color")
    alignment = kwargs.get("alignment")

    if alignment:
        para.alignment = alignment

    # Handle inline bold markers (**text**)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
            run.bold = bold

        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    # Reduce spacing between paragraphs
    para.paragraph_format.space_before = Pt(kwargs.get("space_before", 2))
    para.paragraph_format.space_after = Pt(kwargs.get("space_after", 2))

    return para


def create_docx_from_markdown(markdown_text: str) -> bytes:
    """
    Convert a Markdown-formatted resume into a DOCX file.

    Supports:
        - # H1 (candidate name)
        - ## H2 (section headers)
        - ### H3 (job titles / sub-headers)
        - - Bullet points (with inline **bold**)
        - Regular paragraphs (with inline **bold**)

    Args:
        markdown_text: The resume content in Markdown format.

    Returns:
        The generated DOCX file as bytes.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Narrow margins for a resume
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    lines = markdown_text.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # Skip markdown code fences (```markdown, ```)
        if stripped.startswith("```"):
            continue

        # H1 — Candidate name
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            _add_styled_paragraph(
                doc,
                heading_text,
                "heading1",
                font_size=20,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=0,
                space_after=4,
            )

        # H2 — Section headers
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            heading_text = stripped[3:].strip()
            _add_styled_paragraph(
                doc,
                heading_text,
                "heading2",
                font_size=13,
                bold=True,
                color=(44, 62, 80),  # Dark blue-gray
                space_before=10,
                space_after=3,
            )
            # Add a thin separator line via a bottom border would be complex,
            # so we just add a subtle horizontal rule effect with spacing.

        # H3 — Job titles / sub-headers
        elif stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            _add_styled_paragraph(
                doc,
                heading_text,
                "heading3",
                font_size=11,
                bold=True,
                space_before=6,
                space_after=2,
            )

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            _add_styled_paragraph(
                doc,
                f"•  {bullet_text}",
                "bullet",
                font_size=11,
                space_before=1,
                space_after=1,
            )

        # Horizontal rules (---, ***)
        elif stripped in ("---", "***", "___"):
            continue  # Skip visual separators

        # Regular paragraph
        else:
            _add_styled_paragraph(
                doc,
                stripped,
                "body",
                font_size=11,
                space_before=2,
                space_after=2,
            )

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
